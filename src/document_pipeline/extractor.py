"""
Text extraction from downloaded documents.

Tier 1 — text-based PDFs:  pdfplumber
Tier 2 — DOCX:             python-docx
Tier 3 — scanned PDFs:     PyMuPDF → page images → Sonnet vision
Tier 4 — HTML:             BeautifulSoup

Returns at most MAX_CHARS characters of cleaned text.
"""
from __future__ import annotations

import base64
import io
import logging
import os
from pathlib import Path
from typing import Optional

import requests
import urllib3

logger = logging.getLogger(__name__)

MAX_CHARS = 20_000
_MIN_TEXT_RATIO = 0.01  # if pdfplumber yields < 1% of page area as text → scanned

_SSL_VERIFY = os.environ.get("SSL_VERIFY_DISABLE", "").strip().lower() not in ("1", "true", "yes")
if not _SSL_VERIFY:
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

_OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
_VISION_MODEL = "anthropic/claude-haiku-4.5"


# ── Helpers ───────────────────────────────────────────────────────────────────

def _truncate(text: str) -> str:
    if len(text) <= MAX_CHARS:
        return text
    return text[:MAX_CHARS] + f"\n[… truncated at {MAX_CHARS} chars]"


def _clean_whitespace(text: str) -> str:
    import re
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


# ── PDF via pdfplumber ────────────────────────────────────────────────────────

def _extract_pdf_text(path: Path) -> tuple[str, bool]:
    """Return (text, is_scanned).  is_scanned=True when text is effectively empty."""
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber not installed — skipping PDF text extraction")
        return "", True

    try:
        with pdfplumber.open(path) as pdf:
            pages = pdf.pages[:30]  # cap at 30 pages
            parts: list[str] = []
            total_chars = 0
            for page in pages:
                t = page.extract_text() or ""
                parts.append(t)
                total_chars += len(t)
            text = "\n\n".join(parts)

        if total_chars < 100:
            return text, True  # scanned
        return text, False

    except Exception as e:
        logger.warning(f"pdfplumber error on {path.name}: {e}")
        return "", True


# ── PDF Vision fallback ───────────────────────────────────────────────────────

def _pdf_to_images_b64(path: Path, max_pages: int = 5) -> list[str]:
    """Convert PDF pages to base64 PNG strings via PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning("PyMuPDF not installed — Vision fallback unavailable")
        return []

    images_b64: list[str] = []
    try:
        doc = fitz.open(str(path))
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            mat = fitz.Matrix(1.5, 1.5)  # 1.5× zoom for legibility
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
            png_bytes = pix.tobytes("png")
            images_b64.append(base64.standard_b64encode(png_bytes).decode())
        doc.close()
    except Exception as e:
        logger.warning(f"PyMuPDF error on {path.name}: {e}")

    return images_b64


def _extract_pdf_vision(path: Path) -> str:
    """Send PDF pages as images to Claude Vision via OpenRouter and return extracted text."""
    images_b64 = _pdf_to_images_b64(path)
    if not images_b64:
        return ""

    content = []
    for b64 in images_b64:
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/png;base64,{b64}"},
        })
    content.append({
        "type": "text",
        "text": (
            "This is a scanned procurement document. Extract ALL text you can read, "
            "preserving structure (tables, lists, headers). Focus on: trailer specifications, "
            "quantities, technical requirements, delivery terms. Output raw text only."
        ),
    })

    api_key = (
        os.environ.get("LLM_OPENROUTER_API_KEY")
        or os.environ.get("OPENROUTER_API_KEY")
        or ""
    )
    try:
        resp = requests.post(
            _OPENROUTER_URL,
            json={
                "model": _VISION_MODEL,
                "max_tokens": 4096,
                "messages": [{"role": "user", "content": content}],
            },
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
                "HTTP-Referer": "https://bpw-tender-radar.internal",
                "X-Title": "BPW Defence Tender Radar",
            },
            verify=_SSL_VERIFY,
            timeout=120,
        )
        if resp.status_code != 200:
            logger.warning(f"Vision HTTP {resp.status_code} for {path.name}: {resp.text[:200]}")
            return ""
        return resp.json()["choices"][0]["message"]["content"] or ""
    except Exception as e:
        logger.warning(f"Vision extraction failed for {path.name}: {e}")
        return ""


# ── DOCX ─────────────────────────────────────────────────────────────────────

def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document  # python-docx
    except ImportError:
        logger.warning("python-docx not installed")
        return ""

    try:
        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        logger.warning(f"python-docx error on {path.name}: {e}")
        return ""


# ── HTML ──────────────────────────────────────────────────────────────────────

def _extract_html_text(path: Path) -> str:
    try:
        from bs4 import BeautifulSoup
        html = path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "header", "footer"]):
            tag.decompose()
        return soup.get_text(separator="\n")
    except Exception as e:
        logger.warning(f"HTML extraction error on {path.name}: {e}")
        return ""


# ── Dispatcher ────────────────────────────────────────────────────────────────

def extract_text(
    path: Path,
    fmt: str,
    anthropic_client=None,  # kept for backward-compat; ignored (Vision is self-contained)
) -> str:
    """Extract text from a downloaded document file.

    Falls back to Vision (Haiku via OpenRouter) for scanned PDFs.
    Returns cleaned, truncated text string (empty string on failure).
    """
    fmt = fmt.lower()
    text = ""

    if fmt == "pdf":
        text, is_scanned = _extract_pdf_text(path)
        if is_scanned:
            logger.info(f"Extractor: scanned PDF {path.name} — using Vision fallback")
            text = _extract_pdf_vision(path)

    elif fmt in ("docx", "doc"):
        text = _extract_docx_text(path)

    elif fmt in ("html", "htm"):
        text = _extract_html_text(path)

    elif fmt in ("xlsx", "xls"):
        # openpyxl for xlsx — just dump cell values as text
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            rows = []
            for ws in wb.worksheets[:3]:
                for row in ws.iter_rows(max_row=200, values_only=True):
                    cells = [str(c) for c in row if c is not None and str(c).strip()]
                    if cells:
                        rows.append(" | ".join(cells))
            text = "\n".join(rows)
            wb.close()
        except Exception as e:
            logger.warning(f"xlsx extraction error on {path.name}: {e}")

    else:
        logger.debug(f"Extractor: unsupported format '{fmt}' for {path.name}")

    return _truncate(_clean_whitespace(text))
